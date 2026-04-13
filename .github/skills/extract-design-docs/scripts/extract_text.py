#!/usr/bin/env python3
"""
Extract text content from all .xlsx and .docx files in a directory tree.
Outputs a consolidated text file with clear delimiters per file/sheet.

Usage: python extract_text.py <input_dir> <output_file>
"""

import os
import sys
import subprocess


def ensure_packages():
    """Install required packages if missing."""
    for pkg in ["openpyxl", "python-docx"]:
        try:
            __import__(pkg.replace("-", "_").split(".")[0] if pkg != "python-docx" else "docx")
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])


def extract_xlsx(filepath):
    """Extract text from all sheets of an Excel file."""
    import openpyxl

    results = []
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines = []
            for row in ws.iter_rows(values_only=True):
                cell_values = []
                for cell in row:
                    if cell is not None:
                        cell_values.append(str(cell))
                if cell_values:
                    lines.append("\t".join(cell_values))
            if lines:
                results.append((sheet_name, "\n".join(lines)))
        wb.close()
    except Exception as e:
        results.append(("ERROR", f"Failed to read: {e}"))
    return results


def extract_docx(filepath):
    """Extract text from a Word document."""
    import docx

    results = []
    try:
        doc = docx.Document(filepath)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("\t".join(cells))
        if lines:
            results.append(("Document", "\n".join(lines)))
    except Exception as e:
        results.append(("ERROR", f"Failed to read: {e}"))
    return results


def main():
    if len(sys.argv) < 3:
        print("Usage: python extract_text.py <input_dir> <output_file>")
        sys.exit(1)

    input_dir = sys.argv[1]
    output_file = sys.argv[2]

    ensure_packages()

    file_count = 0
    with open(output_file, "w", encoding="utf-8") as out:
        for root, dirs, files in os.walk(input_dir):
            for fname in sorted(files):
                fpath = os.path.join(root, fname)
                rel_path = os.path.relpath(fpath, input_dir)
                ext = os.path.splitext(fname)[1].lower()

                if ext == ".xlsx":
                    sheets = extract_xlsx(fpath)
                elif ext == ".docx":
                    sheets = extract_docx(fpath)
                else:
                    continue

                file_count += 1
                for sheet_name, content in sheets:
                    out.write("=" * 60 + "\n")
                    out.write(f"FILE: {rel_path}\n")
                    out.write(f"SHEET: {sheet_name}\n")
                    out.write("=" * 60 + "\n")
                    out.write(content + "\n\n")

    print(f"Extracted text from {file_count} files -> {output_file}")


if __name__ == "__main__":
    main()
